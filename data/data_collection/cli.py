import os
from collections import defaultdict
from docx import Document
import requests
from bs4 import BeautifulSoup
from io import BytesIO
from PIL import Image
from google.oauth2 import service_account
from google.cloud import storage


source_data = 'source_data'

if not os.path.exists(source_data):
    os.makedirs(source_data)

def upload_to_gcs(bucket_name, blob_name, content):

    # credentials = service_account.Credentials.from_service_account_file('secrets/acoustic_monitoring_sa.json')

    storage_client = storage.Client(project='acoustic_monitoring_sa')
    bucket = storage_client.bucket(bucket_name)
    
    
    blob = bucket.blob(blob_name)
    blob.upload_from_string(content)
    print(f'Uploaded data to gs://{bucket_name}/{blob_name}')


#extract ebird data
def extract_ebird_data(URL, FILE_PATH):
    if not os.path.exists(FILE_PATH):
        print('downloading data...')
        result = requests.get(URL)
        html = result.text
        with open(FILE_PATH, 'w') as f:
            f.write(html)
            print('written data')
    else:
        print('loading data from file...')
        with open(FILE_PATH, 'r') as f:
            html = f.read()

    soup = BeautifulSoup(html, 'html.parser')
    rows = soup.find("p", class_="u-stack-sm").text
    print(rows)
    # img_url = soup.find('img', class_="Species-media-image")["src"]
    img_url = [img["src"] for img in soup.find_all("img", class_="Species-media-image")]
    print(img_url)
    img_list = []
    for i in img_url:
        img = requests.get(i)
        img_list.append(Image.open(BytesIO(img.content)))
        
    # img = requests.get(img_url)
    # print(img)
    return rows, img_list

def save_doc(text, doc_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(doc_path)

def ebird(spec_name,urls, main_fold):
    for specie, url in zip(spec_name,urls):
       
        path = os.path.join(main_fold, specie)
        os.makedirs(path, exist_ok=True)
        file_name = 'data.html'
        file_path = os.path.join(path, file_name)

        data, images = extract_ebird_data(url, file_path)
        bucket_name = "acoustic_monitoring_project"   
        folder_prefix = "Unprocessed_bird_images"
        folder_prefix2 = "Unprocessed_ebird_text"

        blob_name = f"{folder_prefix2}/{specie.replace(' ', '_')}.docx"

        upload_to_gcs(bucket_name, blob_name, data)
        # upload_to_gcs(bucket_name, f'{folder_prefix2}/{specie}_{i + 1}.jpg', img_p)

        doc_p = os.path.join(path, f'{specie}.docx')
        save_doc(data, doc_p)

        for i, image in enumerate(images):
            img_p = os.path.join(path, f'{specie}_{i + 1}.jpg')
            image.save(img_p)

            upload_to_gcs(bucket_name, f'{folder_prefix}/{specie}_{i + 1}.jpg', img_p)  


#extract peruave data
def extract_peruave_data(URL, FILE_PATH):
    if not os.path.exists(FILE_PATH):
        print('downloading data...')
        result = requests.get(URL)
        html = result.text
        with open(FILE_PATH, 'w') as f:
            f.write(html)
            print('written data')
    else:
        print('loading data from file...')
        with open(FILE_PATH, 'r') as f:
            html = f.read()

    soup = BeautifulSoup(html, 'html.parser')
    rows = soup.find("div", class_="content-column one_half").text
    print(rows)

    return rows


def peru_aves(spec_name,urls, main_fold):
    for specie, url in zip(spec_name,urls):
       
        path = os.path.join(main_fold, specie)
        os.makedirs(path, exist_ok=True)
        file_name = 'data.html'
        file_path = os.path.join(path, file_name)

        data = extract_peruave_data(url, file_path)
        bucket_name = "acoustic_monitoring_project"   
        # folder_prefix = "bird_images"
        folder_prefix2 = "Unprocessed_peru_aves"

        blob_name = f"{folder_prefix2}/{specie.replace(' ', '_')}.docx"
        # upload_to_gcs(bucket_name, f'{folder_prefix}/{specie}_{i + 1}.jpg', img_p) 

        upload_to_gcs(bucket_name, blob_name, data)
        # upload_to_gcs(bucket_name, f'{folder_prefix2}/{specie}_{i + 1}.jpg', img_p) 
        doc_p = os.path.join(path, f'{specie}.docx')
        save_doc(data, doc_p)


#extract - wikipedia
def extract_wiki_data(URL, FILE_PATH):
    if not os.path.exists(FILE_PATH):
        print('downloading data...')
        result = requests.get(URL)
        html = result.content
        with open(FILE_PATH, 'wb') as f:
            f.write(html)
            print('written data')
    else:
        print('loading data from file...')
        with open(FILE_PATH, 'r') as f:
            html = f.read()

    soup = BeautifulSoup(html, 'html.parser')
    # print(soup.prettify())
    rows = soup.find("div", class_="mw-content-ltr mw-parser-output")
    # print(rows)
    rows = rows.find_all('p')
    rows = [row.get_text() for row in rows]
    rows = '\n'.join(rows)
    return rows

def wiki(spec_name,urls, main_fold):
    for specie, url in zip(spec_name,urls):
       
        path = os.path.join(main_fold, specie)
        os.makedirs(path, exist_ok=True)
        file_name = 'data.html'
        file_path = os.path.join(path, file_name)

        data = extract_wiki_data(url, file_path)
        bucket_name = "acoustic_monitoring_project"   
        folder_prefix2 = "Unprocessed_wiki_data"

        blob_name = f"{folder_prefix2}/{specie.replace(' ', '_')}.docx"

        upload_to_gcs(bucket_name, blob_name, data)
        doc_p = os.path.join(path, f'{specie}.docx')
        save_doc(data, doc_p)
    

species = ["Rupicola peruvianus",
           "Andigena hypoglauca",
           "Aulacorhynchus coeruleicinctis",
           "Doliornis sclateri",
           "Pipile cumanensis",
           "Gallinago jamesoni",
           "Tinamus osgoodi",
           "Pionus tumultuosus",
           "Hapalopsittaca melanotis"
           ]

ebird_urls = ["https://ebird.org/species/andcot1",
        "https://ebird.org/species/gybmot1",
        "https://ebird.org/species/blbtou1",
        "https://ebird.org/species/bavcot1",
        "https://ebird.org/species/butpig1",
        "https://ebird.org/species/andsni1",
        "https://ebird.org/species/blatin1",
        "https://ebird.org/species/spfpar1",
        "https://ebird.org/species/blwpar1"
]
mainpath_ebird = os.path.join(source_data, 'unprocessed_ebird')

ebird(species, ebird_urls,mainpath_ebird)


peru_aves_urls = ["https://www.peruaves.org/cotingidae/andean-cock-rock-rupicola-peruvianus/",
        "https://www.peruaves.org/ramphastidae/blue-banded-toucanet-aulacorhynchus-coeruleicinctis/",
        "https://www.peruaves.org/cotingidae/bay-vented-cotinga-doliornis-sclateri/",
        "https://www.peruaves.org/cracidae/blue-throated-piping-guan-pipile-cumanensis/",
        "https://www.peruaves.org/scolopacidae/andean-snipe-gallinago-jamesoni/",
        "https://www.peruaves.org/tinamidae/black-tinamou-tinamus-osgoodi/",
        "https://www.peruaves.org/psittacidae/speckle-faced-parrot-pionus-tumultuosus/",
        "https://www.peruaves.org/psittacidae/black-winged-parrot-hapalopsittaca-melanotis/"
]
mainpath_peruaves = os.path.join(source_data, 'unprocessed_peruaves')

peru_aves(species, peru_aves_urls,mainpath_peruaves)


wiki_urls = ["https://en.wikipedia.org/wiki/Andean_cock-of-the-rock",
        "https://en.wikipedia.org/wiki/Grey-breasted_mountain_toucan",
        "https://en.wikipedia.org/wiki/Blue-banded_toucanet",
        "https://en.wikipedia.org/wiki/Bay-vented_cotinga",
        "https://en.wikipedia.org/wiki/Blue-throated_piping_guan",
        "https://en.wikipedia.org/wiki/Jameson%27s_snipe",
        "https://en.wikipedia.org/wiki/Black_tinamou",
        "https://en.wikipedia.org/wiki/Plum-crowned_parrot",
        "https://en.wikipedia.org/wiki/Black-winged_parrot"
        
]
mainpath_wiki = os.path.join(source_data, 'unprocessed_wiki')
# source_data + 'unprocessed_wiki'

wiki(species, wiki_urls,mainpath_wiki)


#merge all docx and send final data to cloud

script_dir = os.path.dirname(os.path.abspath(__file__))

data_dir = os.path.join(script_dir, "source_data")
subdirs = ["unprocessed_ebird", "unprocessed_peruaves", "unprocessed_wiki"]
output_dir = os.path.join(data_dir, "output_folder")

os.makedirs(output_dir, exist_ok=True)

ebird_path = os.path.join(data_dir, subdirs[0])
if not os.path.exists(ebird_path) or not os.path.isdir(ebird_path):
    print(f"Error: {ebird_path} does not exist or is not a directory.")
    exit(1)

bird_species_dirs = [d for d in os.listdir(ebird_path) if os.path.isdir(os.path.join(ebird_path, d))]


def merge_files(files, output_docx_path, output_txt_path):
    merged_doc = Document()
    merged_text = ""

    for idx, file in enumerate(files):
        if os.path.exists(file):
            doc = Document(file)
            if idx > 0:
                merged_doc.add_paragraph("\n")
                merged_text += "\n"

            for element in doc.paragraphs:
                merged_doc.add_paragraph(element.text)
                merged_text += element.text + "\n"

    merged_doc.save(output_docx_path)

    with open(output_txt_path, "w", encoding="utf-8") as txt_file:
        txt_file.write(merged_text)


for species in bird_species_dirs:
    docx_files = []

    for subdir in subdirs:
        docx_path = os.path.join(data_dir, subdir, species, f"{species}.docx")
        if os.path.exists(docx_path):
            docx_files.append(docx_path)

    if docx_files:
        output_docx_file = os.path.join(output_dir, f"{species}.docx")
        output_txt_file = os.path.join(output_dir, f"{species}.txt")

        merge_files(docx_files, output_docx_file, output_txt_file)
        print(f"Merged files created: {output_docx_file} and {output_txt_file}")




def upload_folder_to_gcs(bucket_name, local_folder, gcs_folder=""):
    # credentials = service_account.Credentials.from_service_account_file('secrets/acoustic_monitoring_sa.json')
    storage_client = storage.Client(project='acoustic_monitoring_sa')
    bucket = storage_client.bucket(bucket_name)

    for root, _, files in os.walk(local_folder):
        for file in files:
            local_path = os.path.join(root, file)
            relative_path = os.path.relpath(local_path, local_folder)
            gcs_path = os.path.join(gcs_folder, relative_path)

            blob = bucket.blob(gcs_path)
            blob.upload_from_filename(local_path)

            print(f"Uploaded {local_path} to gs://{bucket_name}/{gcs_path}")


bucket_name = os.getenv('BUCKET_NAME', 'acoustic_monitoring_project')

local_folder = os.path.join(source_data, 'output_folder')
gcs_folder = "bird_description"

upload_folder_to_gcs(bucket_name, local_folder, gcs_folder)

  

